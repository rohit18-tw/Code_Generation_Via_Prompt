"""Prompt processing module for handling various input types."""

import os
import base64
from pathlib import Path
from typing import List, Dict, Any, Union
from PIL import Image
import PyPDF2
from docx import Document

from config import Config


class PromptProcessor:
    """Handles processing of various input types: text, files, folders, PDFs, images."""
    
    def __init__(self):
        self.config = Config()
    
    def process_input(self, input_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Process input and return structured data for Claude.
        
        Args:
            input_path: Path to file, folder, or direct text
            
        Returns:
            Dict containing processed content and metadata
        """
        input_path = Path(input_path) if isinstance(input_path, str) else input_path
        
        # If it's direct text (not a path)
        if not input_path.exists():
            return self._process_direct_text(str(input_path))
        
        # If it's a file
        if input_path.is_file():
            return self._process_file(input_path)
        
        # If it's a folder
        if input_path.is_dir():
            return self._process_folder(input_path)
        
        raise ValueError(f"Invalid input path: {input_path}")
    
    def _process_direct_text(self, text: str) -> Dict[str, Any]:
        """Process direct text input."""
        return {
            "type": "text",
            "content": text,
            "metadata": {
                "source": "direct_input",
                "length": len(text)
            }
        }
    
    def _process_file(self, file_path: Path) -> Dict[str, Any]:
        """Process a single file based on its extension."""
        extension = file_path.suffix.lower()
        
        if extension in self.config.SUPPORTED_TEXT_EXTENSIONS:
            return self._process_text_file(file_path)
        elif extension in self.config.SUPPORTED_DOC_EXTENSIONS:
            return self._process_document_file(file_path)
        elif extension in self.config.SUPPORTED_IMAGE_EXTENSIONS:
            return self._process_image_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _process_text_file(self, file_path: Path) -> Dict[str, Any]:
        """Process text-based files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                "type": "text_file",
                "content": content,
                "metadata": {
                    "source": str(file_path),
                    "extension": file_path.suffix,
                    "size": len(content),
                    "lines": content.count('\n') + 1
                }
            }
        except Exception as e:
            raise ValueError(f"Error reading text file {file_path}: {e}")
    
    def _process_document_file(self, file_path: Path) -> Dict[str, Any]:
        """Process document files (PDF, DOCX)."""
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._process_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported document type: {extension}")
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF files."""
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            
            return {
                "type": "pdf",
                "content": text.strip(),
                "metadata": {
                    "source": str(file_path),
                    "pages": len(reader.pages),
                    "size": len(text)
                }
            }
        except Exception as e:
            raise ValueError(f"Error reading PDF {file_path}: {e}")
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files."""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "type": "docx",
                "content": text,
                "metadata": {
                    "source": str(file_path),
                    "paragraphs": len(doc.paragraphs),
                    "size": len(text)
                }
            }
        except Exception as e:
            raise ValueError(f"Error reading DOCX {file_path}: {e}")
    
    def _process_image_file(self, file_path: Path) -> Dict[str, Any]:
        """Process image files for Claude vision."""
        try:
            with open(file_path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode('utf-8')
            
            # Get image metadata
            with Image.open(file_path) as img:
                width, height = img.size
                format_type = img.format
            
            return {
                "type": "image",
                "content": image_data,
                "metadata": {
                    "source": str(file_path),
                    "format": format_type,
                    "dimensions": f"{width}x{height}",
                    "size": file_path.stat().st_size
                }
            }
        except Exception as e:
            raise ValueError(f"Error reading image {file_path}: {e}")
    
    def _process_folder(self, folder_path: Path) -> Dict[str, Any]:
        """Process all supported files in a folder."""
        processed_files = []
        
        for file_path in folder_path.rglob('*'):
            if file_path.is_file():
                extension = file_path.suffix.lower()
                if extension in (self.config.SUPPORTED_TEXT_EXTENSIONS | 
                               self.config.SUPPORTED_DOC_EXTENSIONS | 
                               self.config.SUPPORTED_IMAGE_EXTENSIONS):
                    try:
                        file_data = self._process_file(file_path)
                        processed_files.append(file_data)
                    except ValueError as e:
                        print(f"Warning: Skipping file {file_path}: {e}")
        
        # Combine text content from all files
        combined_text = ""
        for file_data in processed_files:
            if file_data["type"] in ["text_file", "pdf", "docx"]:
                combined_text += f"\n\n--- File: {file_data['metadata']['source']} ---\n"
                combined_text += file_data["content"]
        
        return {
            "type": "folder",
            "content": combined_text.strip(),
            "files": processed_files,
            "metadata": {
                "source": str(folder_path),
                "file_count": len(processed_files),
                "total_size": sum(f["metadata"].get("size", 0) for f in processed_files)
            }
        }
